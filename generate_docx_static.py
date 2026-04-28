from pathlib import Path
from datetime import datetime

from docx import Document

from app import app


def _resolve_schema(schema: dict, components: dict) -> dict:
    if not schema:
        return {}
    ref = schema.get("$ref")
    if not ref:
        return schema
    prefix = "#/components/schemas/"
    if ref.startswith(prefix):
        schema_name = ref[len(prefix) :]
        return components.get("schemas", {}).get(schema_name, {})
    return schema


def _add_schema_lines(doc: Document, schema: dict, indent: int = 0) -> None:
    if not schema:
        return

    schema_type = schema.get("type", "object")
    prefix = "  " * indent
    doc.add_paragraph(f"{prefix}type: {schema_type}", style="List Bullet")

    if schema_type == "object":
        properties = schema.get("properties", {})
        required_fields = set(schema.get("required", []))
        for prop_name, prop_schema in properties.items():
            prop_type = prop_schema.get("type", "object")
            required = "required" if prop_name in required_fields else "optional"
            doc.add_paragraph(
                f"{prefix}{prop_name}: {prop_type} ({required})", style="List Bullet"
            )
            if prop_schema.get("type") in {"object", "array"}:
                _add_schema_lines(doc, prop_schema, indent + 1)

    elif schema_type == "array":
        items = schema.get("items", {})
        item_type = items.get("type", "object")
        doc.add_paragraph(f"{prefix}items: {item_type}", style="List Bullet")
        if item_type in {"object", "array"}:
            _add_schema_lines(doc, items, indent + 1)


def _format_parameter(param: dict) -> str:
    name = param.get("name", "unknown")
    where = param.get("in", "unknown")
    required = param.get("required", False)
    param_type = param.get("schema", {}).get("type", "object")
    return f"{name} ({where}, type={param_type}, required={required})"


def main() -> None:
    with app.test_client() as client:
        response = client.get("/openapi.json")
        if response.status_code != 200:
            raise RuntimeError(
                f"Failed to get OpenAPI spec. Status code: {response.status_code}"
            )
        spec = response.get_json()
        if not spec:
            raise RuntimeError("OpenAPI spec is empty or invalid JSON.")

    info = spec.get("info", {})
    doc = Document()
    doc.add_heading(info.get("title", "API Documentation"), 0)
    doc.add_paragraph(f"Version: {info.get('version', 'N/A')}")
    if info.get("description"):
        doc.add_paragraph(info["description"])

    components = spec.get("components", {})
    paths = spec.get("paths", {})
    for path, operations in paths.items():
        doc.add_heading(path, level=1)
        for method, details in operations.items():
            doc.add_heading(method.upper(), level=2)
            if details.get("summary"):
                doc.add_paragraph(f"Summary: {details['summary']}")
            if details.get("description"):
                doc.add_paragraph(f"Description: {details['description']}")

            parameters = details.get("parameters", [])
            if parameters:
                doc.add_paragraph("Parameters:")
                for param in parameters:
                    doc.add_paragraph(_format_parameter(param), style="List Bullet")

            request_body = details.get("requestBody", {})
            if request_body:
                required = request_body.get("required", False)
                doc.add_paragraph(f"Request Body (required={required}):")
                media_types = request_body.get("content", {})
                for media_type, media_meta in media_types.items():
                    doc.add_paragraph(f"Media Type: {media_type}", style="List Bullet")
                    body_schema = _resolve_schema(media_meta.get("schema", {}), components)
                    if body_schema:
                        doc.add_paragraph("Request Body Schema:")
                        _add_schema_lines(doc, body_schema, indent=1)

            responses = details.get("responses", {})
            if responses:
                doc.add_paragraph("Responses:")
                for status_code, response_meta in responses.items():
                    description = response_meta.get("description", "")
                    doc.add_paragraph(
                        f"{status_code}: {description}", style="List Bullet"
                    )
                    content = response_meta.get("content", {})
                    for media_type, media_meta in content.items():
                        doc.add_paragraph(
                            f"Response Media Type: {media_type}", style="List Bullet"
                        )
                        response_schema = _resolve_schema(
                            media_meta.get("schema", {}), components
                        )
                        if response_schema:
                            doc.add_paragraph("Response Body Schema:")
                            _add_schema_lines(doc, response_schema, indent=1)

    output_dir = Path("docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "api-docs.docx"
    try:
        doc.save(output_file)
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = output_dir / f"api-docs-{timestamp}.docx"
        doc.save(output_file)
    print(f"Generated static DOCX at: {output_file.resolve()}")


if __name__ == "__main__":
    main()
